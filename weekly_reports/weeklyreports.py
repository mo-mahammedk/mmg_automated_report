#!/usr/bin/python
from docx import Document
from docx.shared import Cm, Pt
from friday import friday, worded
import getpass
# import markdown

document = Document()
style = document.styles['Normal']
style.paragraph_format.line_spacing = 1

# Title of Document 
# document.add_heading('Monsoon2 Weekly Service Report', 0) 

# Title of Document
title = document.add_paragraph().add_run(
    'Monsoon2 Weekly Service Report')
title.font.name = 'Arial'
title.font.size = Pt(14)
title.bold = True

# Report Information
report_info = f'Report for week ending: {worded}\n(Monsoon report runs from 12:00 Friday – 12:00 Friday, MASS 09:00 Thursday – 09:00 Thursday).'
date,time = report_info.split('\n', 1)
wk_ending = document.add_paragraph().add_run(
    date)
wk_ending.font.name = 'Calibri (Body)'
wk_ending.font.size = Pt(13)
wk_ending.bold = True
report_runs = document.add_paragraph().add_run(
    time)
report_runs.font.name = 'Calibri (Body)'
report_runs.font.size = Pt(10)

# Emtpy line
document.add_paragraph().add_run(
    '')

# Planned Maintenance
planned = document.add_paragraph().add_run(
    'Planned Maintenance')
planned.font.name = 'Calibri (Body)'
planned.font.size = Pt(11)
planned.bold = True
planned.underline = True
# Insert from docs any planned maintence
with open('planned.txt', 'r') as f:
   lines = f.read()
    
remove = str(lines).replace('[', '').replace(']', '').replace("'", "").replace(',', '')
plannedlist = document.add_paragraph().add_run(
    f'{remove}')

plannedlist.font.name = 'Calibri (Body)'
plannedlist.font.size = Pt(11)


# Unplanned Maintenance
unplanned = document.add_paragraph().add_run(
    'Unplanned Maintenance')
unplanned.font.name = 'Calibri (Body)'
unplanned.font.size = Pt(11)
unplanned.bold = True
unplanned.underline = True
# Insert from docs any unplanned maintence
with open('unplanned.txt', 'r') as f2:
   lines2 = f2.read()
    
remove2 = str(lines2).replace('[', '').replace(']', '').replace("'", "").replace(',', '')
unplannedlist = document.add_paragraph().add_run(
    f'{remove2}')

unplannedlist.font.name = 'Calibri (Body)'
unplannedlist.font.size = Pt(11)


# Upcoming Maintenance
upcoming = document.add_paragraph().add_run(
    'Upcoming Maintenance')
upcoming.font.name = 'Calibri (Body)'
upcoming.font.size = Pt(11)
upcoming.bold = True
upcoming.underline = True
# Insert from docs any Upcoming maintence
with open('upcoming.txt', 'r') as f3:
   lines3 = f3.read()
    
remove3 = str(lines3).replace('[', '').replace(']', '').replace("'", "").replace(',', '')
upcominglist = document.add_paragraph().add_run(
    f'{remove3}')

upcominglist.font.name = 'Calibri (Body)'
upcominglist.font.size = Pt(11)


# Emtpy line
document.add_paragraph().add_run(
    '')

# Service Summary
summary = document.add_paragraph().add_run(
    'Service Summary')
summary.font.name = 'Calibri (Body)'
summary.font.size = Pt(11)
summary.bold = True
summary.underline = True

# Service bullet points - traffic lights
rag_status = ('Red', 'Amber', 'Green')


monsoon_input = input('Select a RAG status for Monsoon - Red, Amber or Green: ')
if monsoon_input not in rag_status:
    raise Exception('Option not available')

mass_input = input('Select a RAG status for MASS - Red, Amber or Green: ')
if mass_input not in rag_status:
    raise Exception('Option not available')

    
monsoon = document.add_paragraph().add_run(
    f'• Monsoon2 service is {monsoon_input}. \n• MASS service is {mass_input}.')
monsoon.font.name = 'Calibri (Body)'
monsoon.font.size = Pt(11)

# document.add_paragraph("Monsoon2 service is ", style="List Bullet")
# document.add_paragraph("MASS service is ", style="List Bullet")

# Emtpy line
document.add_paragraph().add_run(
    '')

# Signature


if getpass.getuser() == 'mkassam':
    ending = document.add_paragraph().add_run(
        'Mahammed Kassam \nMonsoon Collaboration Team Member')
    ending.font.name = 'Calibri (Body)'
    ending.font.size = Pt(11)
elif getpass.getuser() == 'itrm':
    ending = document.add_paragraph().add_run(
        'Roger Milton \nMonsoon Technical Lead')
    ending.font.name = 'Calibri (Body)'
    ending.font.size = Pt(11)
elif getpass.getuser() == 'mking':
    ending = document.add_paragraph().add_run(
        'Michael King \nMonsoon Collaboration Team Member')
    ending.font.name = 'Calibri (Body)'
    ending.font.size = Pt(11)
# document.add_page_break()

document.save(f'Monsoon2 Weekly Report - {friday}.docx')

print('---------------------------------------------')
print(f'Weekly report for {friday} has been created')
print('---------------------------------------------')